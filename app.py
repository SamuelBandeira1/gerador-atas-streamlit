"""
Gerador de Atas - Secretário Executivo de IA
Grupo Aço Cearense — S&OE / S&OP
"""

import io
import json
import re
import unicodedata
import streamlit as st
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from datetime import datetime

# ─────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────
MEETING_TYPES = {
    "Supply Aço":  "Reunião de Alinhamento S&OE Aço (Supply)",
    "Supply SIN":  "Reunião de Alinhamento S&OE SIN (Supply)",
    "Demanda":     "Reunião de Alinhamento S&OE (Demanda)",
}

STATUS_OPTIONS = ["A iniciar", "Em andamento", "Em atraso", "Concluído"]

MODEL_NAME = "gemini-2.5-pro-exp-03-25"  # default

MODEL_OPTIONS = [
    "gemini-2.5-pro-exp-03-25",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-pro",
]

# ─────────────────────────────────────────────
# DICIONÁRIO DE CORREÇÃO DE TRANSCRIÇÃO
# ─────────────────────────────────────────────
# Mapeia erros comuns de reconhecimento de voz → termo correto.
# Padrões são case-insensitive; substituição preserva o caso do original quando possível.
CORRECTION_DICT = [
    # Produtos / Famílias
    (r"\bca\s*[\-–]?\s*50\b",                   "CA-50"),
    (r"\bca\s*[\-–]?\s*60\b",                   "CA-60"),
    (r"\bvergalhão\b",                           "Vergalhão"),
    (r"\bfio\s*máquina\b",                        "Fio-máquina"),
    (r"\bfio\s*maquina\b",                        "Fio-máquina"),
    (r"\bfio[\-\s]máq\b",                         "Fio-máquina"),
    (r"\bchapa\s*plana\b",                        "Chapa Plana"),
    (r"\bchapa\s*grossa\b",                       "Chapa Grossa"),
    (r"\bbobina\s*reduzida\b",                    "Bobina Reduzida"),
    (r"\bgalvamax\b",                             "Galvamax"),
    (r"\bgalvanizado\b",                          "Galvanizado"),
    (r"\bgalvanizada\b",                          "Galvanizada"),
    (r"\binox\b",                                 "Inox"),
    (r"\bperfil\b",                               "Perfil"),
    (r"\btelha\b",                                "Telha"),
    (r"\btela\s*coluna\b",                        "Tela Coluna"),
    (r"\btela\b",                                 "Tela"),
    (r"\btubo\b",                                 "Tubo"),
    (r"\bcantoneira\b",                           "Cantoneira"),
    (r"\bfamília\s*de\s*longos\b",               "Família de Longos"),
    (r"\bplanos\b",                               "Planos"),
    (r"\btarugo\b",                               "Tarugo"),
    (r"\bcarboneto\b",                            "Carboneto"),
    (r"\bcalcário\b",                             "Calcário"),
    (r"\bcalc[aá]cio\b",                           "Cálcio"),
    (r"\bcalc[ií]o\b",                             "Cálcio"),
    # Empresas / Unidades
    (r"\bsino\s*bra[sz]\b",                       "Sinobras"),
    (r"\bsinobras\b",                             "Sinobras"),
    (r"\baço\s+cearense\b",                       "Aço Cearense"),
    (r"\bgrupo\s+aço\s+cearense\b",              "Grupo Aço Cearense"),
    (r"\bmaia\b",                                 "Maia"),
    (r"\bkmf\b",                                  "KMF"),
    (r"\bunicon\b",                               "Unicon"),
    # Processos / Sistemas
    (r"\bs[\s&]?[&e]?[\s]?o[\s]?e\b",           "S&OE"),
    (r"\bsoe\b",                                  "S&OE"),
    (r"\bs[\s&]?[&e]?[\s]?o[\s]?p\b",           "S&OP"),
    (r"\bsop\b",                                  "S&OP"),
    (r"\bpcp\b",                                  "PCP"),
    (r"\bcic\b",                                  "CIC"),
    (r"\bsap\b",                                  "SAP"),
    (r"\bibp\b",                                  "IBP"),
    (r"\bfifo\b",                                 "FIFO"),
    (r"\bcrm\b",                                  "CRM"),
    (r"\bsku[s]?\b",                              "SKU"),
    (r"\bintercompany\b",                         "Intercompany"),
    (r"\bcooispi\b",                              "COOISPI"),
    # Logística
    (r"\bcabotagem\b",                            "Cabotagem"),
    (r"\bfracionamento\b",                        "Fracionamento"),
    (r"\broteirização\b",                         "Roteir ização"),
    (r"\broteirizaçao\b",                         "Roteiri zação"),
    (r"\bfob\b",                                  "FOB"),
    (r"\banti\s*dumping\b",                       "Antidumping"),
    (r"\bantidamping\b",                          "Antidumping"),
    # Indicadores
    (r"\bclose\s*the\s*gap\b",                   "Close-the-Gap"),
    (r"\bm\s*\+\s*1\b",                           "M+1"),
    (r"\bbudget\b",                               "Budget"),
    (r"\bbacklog\b",                              "Backlog"),
    # CA-50 / CA-60: erros fonéticos ("Cássia 50", "CASSIA 50", "Cassio 60")
    (r"\bcass?[iíy]?[ao]?\s*[\-–]?\s*50\b",      "CA-50"),
    (r"\bcass?[iíy]?[ao]?\s*[\-–]?\s*60\b",      "CA-60"),
    (r"\bCA\s+50\b",                              "CA-50"),
    (r"\bCA\s+60\b",                              "CA-60"),
    # COOISPI: "coísp", "coíspe", "coíspi", "coíspia"
    (r"\bco[ií]sp[iea]?\b",                       "COOISPI"),
    (r"\bco[ií]spi\b",                            "COOISPI"),
    # Dashboard
    (r"\bdash(?:board)?\b",                       "Dashboard"),
    # "bering dutor" → provavelmente Vergalhão
    (r"\bbering\s+dutor\b",                       "Vergalhão"),
    # Alto-Forno variações
    (r"\balto[\s\-]*forno\b",                     "Alto-Forno"),
    # Trefilaria / Trefila
    (r"\btrefil[a-z]*\b",                         "Trefilaria"),
    # Fio-máquina fonético
    (r"\bfio\s+de\s+m[áa]quina\b",               "Fio-máquina"),
]


def preprocess_transcription(text: str) -> str:
    """Apply correction dictionary to fix common speech-recognition errors."""
    for pattern, replacement in CORRECTION_DICT:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return text

SYSTEM_PROMPT = """
Você é um Secretário Executivo especializado em S&OE (Sales & Operations Execution) do Grupo Aço Cearense.
Sua tarefa é analisar a transcrição bruta de uma reunião e extrair as informações estruturadas.

Retorne SOMENTE um JSON válido (sem markdown, sem texto adicional) com a seguinte estrutura:

{
  "panorama_geral": "resumo do panorama geral do ciclo",
  "plano_do_mes": "informações sobre o plano do mês (volume, referências)",
  "realizado_acumulado": "o que foi realizado até o momento da reunião (volume, percentual)",
  "ritmo_vs_meta": "análise do ritmo atual vs meta estabelecida",
  "principais_desvios": "lista dos principais desvios identificados",
  "fatores_estruturais": "fatores estruturais identificados que impactam o resultado",
  "riscos_prospectivos": "riscos futuros identificados (volume, margem, operacional, estratégico)",
  "dinamica_comercial": "dinâmica comercial do período",
  "performance_linha_familia": "performance por linha/família de produtos",
  "performance_regional": "performance por região",
  "decisoes_tomadas": "decisões explícitas tomadas no fórum (lista)",
  "participantes": [
    {"area": "nome da área", "nomes": "nomes dos participantes"}
  ],
  "acoes": [
    {
      "id": 1,
      "acao": "descrição clara e completa da ação a ser executada",
      "contexto": "por que esta ação existe, o que motivou",
      "responsavel": "nome(s) do(s) responsável(is)",
      "status": "A iniciar | Em andamento | Em atraso | Concluído",
      "prazo": "data no formato DD/MM/AAAA"
    }
  ]
}

REGRAS IMPORTANTES:
1. Se uma seção não foi discutida na reunião, preencha com: "Não discutido nesta sessão"
2. Identifique os participantes agrupados por área (Planejamento Integrado, Comercial, PCP, Customer Service, Logística, CIC, Suprimentos, etc.)
3. Extraia TODAS as ações e tarefas mencionadas, mesmo que implícitas
4. Mantenha o Portuguese (pt-BR) formal e corporativo
5. Retorne APENAS o JSON, sem qualquer texto antes ou depois
"""

# ─────────────────────────────────────────────
# HELPER: Word document utilities
# ─────────────────────────────────────────────

def set_cell_background(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_borders(table, color: str = "BFBFBF", size: int = 4):
    """Apply subtle borders to all cells of a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_paragraph(doc, text: str, bold: bool = False, font_size: int = 11,
                  space_before: int = 0, space_after: int = 4,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    """Add a styled paragraph to the document."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(font_size * 1.35)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_section_heading(doc, number: int, title: str):
    """Add a numbered bold section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(15)
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def to_str(value, fallback: str = "Não discutido nesta sessão") -> str:
    """Safely convert any Gemini field value to a plain string."""
    if value is None:
        return fallback
    if isinstance(value, str):
        return value.strip() or fallback
    if isinstance(value, list):
        parts = []
        for item in value:
            if isinstance(item, dict):
                parts.append("; ".join(f"{k}: {v}" for k, v in item.items()))
            else:
                parts.append(str(item))
        return "\n".join(parts) or fallback
    if isinstance(value, dict):
        return "\n".join(f"{k}: {v}" for k, v in value.items()) or fallback
    return str(value)


def add_section_text(doc, text: str):
    """Add body text under a section, handling bullet-like content."""
    if not text or not isinstance(text, str) or text.strip() == "":
        text = "Não discutido nesta sessão"
    lines = text.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Detect bullet lines (starts with -, •, *, or numbered)
        is_bullet = bool(re.match(r"^[-•*\u2022]|^\d+\.", line))
        clean_line = re.sub(r"^[-•*\u2022]\s*|\d+\.\s*", "", line) if is_bullet else line
        p = doc.add_paragraph(style="List Paragraph" if is_bullet else "Normal")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(15)
        if is_bullet:
            p.paragraph_format.left_indent = Cm(0.63)
        run = p.add_run(clean_line)
        run.font.name = "Calibri"
        run.font.size = Pt(11)


# ─────────────────────────────────────────────
# CORE: Generate Word document
# ─────────────────────────────────────────────

def generate_docx(data: dict, meeting_type: str, week: str, date: str, location: str) -> bytes:
    """Build the DOCX from the structured Gemini output."""
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Title ──
    meeting_label = MEETING_TYPES.get(meeting_type, meeting_type)
    add_paragraph(doc, f"{meeting_label} – {week}",
                  bold=True, font_size=14, space_after=4)

    # ── Date & Location ──
    add_paragraph(doc, f"Data: {date} | Local: {location}",
                  bold=True, font_size=11, space_after=6)

    # ── Participants label ──
    add_paragraph(doc, "Participantes:", bold=True, font_size=11, space_after=4)

    # ── Participants table ──
    participantes = data.get("participantes", [])
    if participantes:
        tbl = doc.add_table(rows=len(participantes), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_borders(tbl, color="BFBFBF", size=4)
        for i, part in enumerate(participantes):
            row = tbl.rows[i]
            # Area cell
            area_cell = row.cells[0]
            area_cell.width = Cm(4.5)
            area_p = area_cell.paragraphs[0]
            area_run = area_p.add_run(part.get("area", ""))
            area_run.bold = True
            area_run.font.name = "Calibri"
            area_run.font.size = Pt(10)
            area_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Names cell
            names_cell = row.cells[1]
            names_p = names_cell.paragraphs[0]
            names_run = names_p.add_run(part.get("nomes", ""))
            names_run.font.name = "Calibri"
            names_run.font.size = Pt(10)
            names_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        doc.add_paragraph()  # spacer

    # ── Numbered sections ──
    sections = [
        (1,  "Panorama Geral do Ciclo",          to_str(data.get("panorama_geral"))),
        (2,  "Plano do Mês",                     to_str(data.get("plano_do_mes"))),
        (3,  "Realizado Acumulado",               to_str(data.get("realizado_acumulado"))),
        (4,  "Ritmo vs Meta",                    to_str(data.get("ritmo_vs_meta"))),
        (5,  "Principais Desvios",                to_str(data.get("principais_desvios"))),
        (6,  "Fatores Estruturais Identificados", to_str(data.get("fatores_estruturais"))),
        (7,  "Riscos Prospectivos",               to_str(data.get("riscos_prospectivos"))),
        (8,  "Dinâmica Comercial do Período",     to_str(data.get("dinamica_comercial"))),
        (9,  "Performance por Linha / Família",   to_str(data.get("performance_linha_familia"))),
        (10, "Performance Regional",              to_str(data.get("performance_regional"))),
        (11, "Decisões Tomadas",                  to_str(data.get("decisoes_tomadas"))),
    ]

    for num, title, content in sections:
        add_section_heading(doc, num, title)
        add_section_text(doc, content)

    # ── Section 12: Action Plan table ──
    add_section_heading(doc, 12, "Saídas/Ações")

    acoes = data.get("acoes", [])
    if acoes:
        col_widths = [Cm(1.0), Cm(6.0), Cm(5.0), Cm(2.5), Cm(2.0), Cm(2.0)]
        headers = ["ID", "Saída/Ação", "Contexto/Referência", "Responsável", "Status", "Prazo"]

        tbl = doc.add_table(rows=1 + len(acoes), cols=6)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_table_borders(tbl, color="BFBFBF", size=4)

        # Header row
        hdr = tbl.rows[0]
        for j, (h, w) in enumerate(zip(headers, col_widths)):
            cell = hdr.cells[j]
            cell.width = w
            set_cell_background(cell, "1F3864")   # dark navy
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Data rows
        for i, acao in enumerate(acoes):
            row = tbl.rows[i + 1]
            # Alternate row shading
            bg = "EBF0FA" if i % 2 == 0 else "FFFFFF"
            values = [
                str(acao.get("id", i + 1)),
                acao.get("acao", ""),
                acao.get("contexto", ""),
                acao.get("responsavel", ""),
                acao.get("status", "A iniciar"),
                acao.get("prazo", ""),
            ]
            for j, (val, w) in enumerate(zip(values, col_widths)):
                cell = row.cells[j]
                cell.width = w
                set_cell_background(cell, bg)
                p = cell.paragraphs[0]
                if j == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    else:
        add_section_text(doc, "Não discutido nesta sessão")

    # ── Save to bytes buffer ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────
# CORE: Call Gemini API
# ─────────────────────────────────────────────

def call_gemini(api_key: str, transcription: str, meeting_type: str, model: str = MODEL_NAME) -> dict:
    """Send transcription to Gemini and return structured dict."""
    client = genai.Client(api_key=api_key)
    # Apply terminology correction before sending to the model
    transcription = preprocess_transcription(transcription)
    user_prompt = f"""
Tipo de reunião: {meeting_type}

Transcrição da reunião:
---
{transcription}
---

Extraia todas as informações e retorne SOMENTE o JSON estruturado, sem markdown.
"""
    response = client.models.generate_content(
        model=model,
        contents=user_prompt,
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_PROMPT,
            temperature=0.2,
        ),
    )
    raw = response.text.strip()
    # Strip markdown code fences if present
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


# ─────────────────────────────────────────────
# CORE: Consolidate split transcription files
# ─────────────────────────────────────────────

def sort_key_from_filename(filename: str) -> int:
    """Extract numeric part index from filenames like 'xxx_parte_02_de_07.txt'."""
    m = re.search(r'parte[_\s]*(\d+)', filename, re.IGNORECASE)
    return int(m.group(1)) if m else 0


def consolidate_uploaded_files(uploaded_files) -> tuple[str, list[dict]]:
    """Sort files by part number and concatenate their text content.
    Returns (consolidated_text, file_metadata_list)."""
    sorted_files = sorted(uploaded_files, key=lambda f: sort_key_from_filename(f.name))
    parts = []
    metadata = []
    for f in sorted_files:
        try:
            content = f.read().decode("utf-8", errors="replace")
        except Exception:
            content = f.read().decode("latin-1", errors="replace")
        parts.append(content)
        metadata.append({"nome": f.name, "tamanho": f.size, "chars": len(content)})
    consolidated = "\n\n".join(parts)
    return consolidated, metadata


# ─────────────────────────────────────────────
# STREAMLIT UI
# ─────────────────────────────────────────────

def main():
    st.set_page_config(
        page_title="Gerador de Atas — S&OE | Aço Cearense",
        page_icon="📋",
        layout="wide",
    )

    # ── CSS ──
    st.markdown("""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');
        html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
        .main-title {
            font-size: 2rem; font-weight: 700;
            background: linear-gradient(135deg, #1F3864, #2E6DB4);
            -webkit-background-clip: text; -webkit-text-fill-color: transparent;
            margin-bottom: 0.2rem;
        }
        .subtitle {
            color: #6c757d; font-size: 0.95rem; margin-bottom: 1.5rem;
        }
        .stButton > button {
            background: linear-gradient(135deg, #1F3864, #2E6DB4);
            color: white; border: none; border-radius: 8px;
            padding: 0.6rem 1.4rem; font-weight: 600; font-size: 1rem;
            transition: all 0.3s ease; width: 100%;
        }
        .stButton > button:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 20px rgba(31,56,100,0.4);
        }
        .info-card {
            background: #f0f4ff;
            border-left: 4px solid #2E6DB4;
            border-radius: 6px; padding: 0.8rem 1rem;
            margin-bottom: 1rem; font-size: 0.9rem;
        }
        .section-header {
            color: #1F3864; font-weight: 600;
            font-size: 1.05rem; margin-top: 1.2rem;
        }
        [data-testid="stSidebar"] {
            background: linear-gradient(180deg, #0f1f3d 0%, #1F3864 100%);
        }
        [data-testid="stSidebar"] * { color: #e8eeff !important; }
        [data-testid="stSidebar"] input { background: #2a3f6b !important; color: white !important; }
        [data-testid="stSidebar"] label { color: #c5d0f0 !important; font-size: 0.85rem !important; }
        [data-testid="stSidebar"] h2 { color: white !important; font-size: 1rem !important; }
    </style>
    """, unsafe_allow_html=True)

    # ── Sidebar ──
    with st.sidebar:
        st.image("https://upload.wikimedia.org/wikipedia/commons/3/3c/Logo_A%C3%A7o_Cearense.svg",
                 use_container_width=True, output_format="auto")
    with st.sidebar:
        st.markdown("## ⚙️ Configurações")
        st.markdown("---")

        # API Key
        api_key = st.text_input(
            "🔑 Gemini API Key",
            type="password",
            placeholder="AIza...",
            help="Sua chave da API do Google Gemini",
        )

        st.markdown("### 📄 Dados da Reunião")
        meeting_type = st.selectbox(
            "Tipo de Reunião",
            options=list(MEETING_TYPES.keys()),
            help="Selecione o tipo de fórum S&OE",
        )

        week = st.text_input("Semana (ex: W13)", value="W13", max_chars=6)

        date = st.text_input(
            "Data da Reunião",
            value=datetime.today().strftime("%d/%m/%Y"),
            help="Formato: DD/MM/AAAA",
        )

        location = st.text_input(
            "Local",
            value="Sala do 14º e Teams",
            help="Ex: Sala do 14º e Teams | Via Teams",
        )

        st.markdown("### 🤖 Modelo Gemini")

        # Dynamic model listing
        if "available_models" not in st.session_state:
            st.session_state.available_models = MODEL_OPTIONS

        if st.button("🔍 Listar Modelos Disponíveis", use_container_width=True):
            if not api_key:
                st.warning("Insira a API Key primeiro.")
            else:
                try:
                    _client = genai.Client(api_key=api_key)
                    _models = [
                        m.name.replace("models/", "")
                        for m in _client.models.list()
                        if "generateContent" in (m.supported_actions or [])
                        and "gemini" in m.name
                    ]
                    if _models:
                        st.session_state.available_models = _models
                        st.success(f"✅ {len(_models)} modelo(s) encontrado(s)!")
                    else:
                        st.warning("Nenhum modelo Gemini encontrado.")
                except Exception as ex:
                    st.error(f"Erro: {ex}")

        selected_model = st.selectbox(
            "Modelo",
            options=st.session_state.available_models,
            index=0,
            help="Clique em 'Listar Modelos' para ver os disponíveis para sua chave.",
        )

        st.markdown("---")
        st.caption("📌 Grupo Aço Cearense — S&OE")
        st.caption("🤖 Powered by Google Gemini")

    # ── Main Content ──
    st.markdown('<p class="main-title">📋 Gerador de Atas — Secretário de IA</p>', unsafe_allow_html=True)
    st.markdown('<p class="subtitle">Grupo Aço Cearense · S&OE / S&OP · Processamento Inteligente de Reuniões</p>', unsafe_allow_html=True)

    # ── Input Tabs ──
    tab_files, tab_text = st.tabs([
        "📂 Arquivos Divididos (Partes)",
        "📋 Colar Transcrição Manual",
    ])

    transcription = ""
    file_metadata = []

    with tab_files:
        st.markdown('<div class="info-card">📂 <b>Modo de arquivos:</b> Selecione todos os arquivos <code>.txt</code> das partes da transcrição. O app vai ordenar automaticamente por número de parte e consolidar tudo antes de enviar para a IA.</div>', unsafe_allow_html=True)
        uploaded_files = st.file_uploader(
            "Selecione os arquivos de transcrição (.txt)",
            type=["txt"],
            accept_multiple_files=True,
            label_visibility="collapsed",
        )
        if uploaded_files:
            consolidated, file_metadata = consolidate_uploaded_files(uploaded_files)
            transcription = consolidated
            st.markdown(f"**{len(uploaded_files)} arquivo(s) carregado(s) — {len(transcription):,} caracteres totais**")
            cols = st.columns(min(len(file_metadata), 4))
            for i, meta in enumerate(file_metadata):
                with cols[i % len(cols)]:
                    st.metric(
                        label=meta["nome"].split("_parte_")[-1].replace(".txt", "") if "_parte_" in meta["nome"] else meta["nome"][:20],
                        value=f"{meta['chars']:,} chars",
                        delta=f"{meta['tamanho'] / 1024:.1f} KB",
                    )

    with tab_text:
        st.markdown('<div class="info-card">📋 <b>Modo manual:</b> Cole a transcrição bruta abaixo (texto de áudio, anotações, etc.).</div>', unsafe_allow_html=True)
        manual_text = st.text_area(
            label="Transcrição manual",
            height=350,
            placeholder="Cole aqui a transcrição completa da reunião...",
            label_visibility="collapsed",
        )
        if manual_text.strip():
            transcription = manual_text

    col1, col2, col3 = st.columns([2, 1, 2])
    with col2:
        process_btn = st.button("🤖 Processar com IA e Gerar Ata", use_container_width=True)

    st.markdown("---")

    # ── Processing ──
    if process_btn:
        # Validations
        if not api_key:
            st.error("❌ Insira a Gemini API Key na barra lateral antes de processar.")
            st.stop()
        if not transcription or len(transcription.strip()) < 50:
            st.error("❌ Nenhuma transcrição encontrada. Carregue os arquivos .txt ou cole o texto manualmente.")
            st.stop()
        if not week.strip():
            st.error("❌ Informe o número da semana (ex: W13).")
            st.stop()
        if file_metadata:
            st.info(f"📂 Consolidando {len(file_metadata)} parte(s) — {len(transcription):,} caracteres no total.")

        # Process
        result_placeholder = st.empty()
        with st.spinner(f"🧠 Analisando com {selected_model}... Isso pode levar até 60 segundos."):
            try:
                structured = call_gemini(api_key, transcription, meeting_type, selected_model)
            except json.JSONDecodeError as e:
                st.error(f"❌ Erro ao interpretar a resposta da IA. Tente novamente.\nDetalhe: {e}")
                st.stop()
            except Exception as e:
                st.error(f"❌ Erro na API Gemini: {e}")
                st.stop()

        st.success("✅ Análise concluída! Gerando documento Word...")

        # Generate DOCX
        try:
            docx_bytes = generate_docx(structured, meeting_type, week, date, location)
        except Exception as e:
            st.error(f"❌ Erro ao gerar o documento Word: {e}")
            st.stop()

        # Show preview
        st.markdown("### 📊 Resumo Extraído pela IA")

        with st.expander("👥 Participantes", expanded=False):
            participantes = structured.get("participantes", [])
            if participantes:
                df = pd.DataFrame(participantes)
                df.columns = ["Área", "Participantes"]
                st.dataframe(df, use_container_width=True, hide_index=True)
            else:
                st.info("Nenhum participante identificado.")

        with st.expander("📋 Indicadores S&OE", expanded=True):
            col_a, col_b = st.columns(2)
            with col_a:
                st.markdown("**🔍 Panorama Geral**")
                st.info(structured.get("panorama_geral", "—")[:500])
                st.markdown("**📅 Plano do Mês**")
                st.info(structured.get("plano_do_mes", "—")[:300])
                st.markdown("**✅ Realizado Acumulado**")
                st.info(structured.get("realizado_acumulado", "—")[:300])
            with col_b:
                st.markdown("**📈 Ritmo vs Meta**")
                st.info(structured.get("ritmo_vs_meta", "—")[:300])
                st.markdown("**⚠️ Principais Desvios**")
                st.warning(structured.get("principais_desvios", "—")[:400])

        with st.expander("🚨 Gargalos e Riscos", expanded=False):
            st.markdown("**🔧 Fatores Estruturais**")
            st.info(structured.get("fatores_estruturais", "—"))
            st.markdown("**⚡ Riscos Prospectivos**")
            st.warning(structured.get("riscos_prospectivos", "—"))

        with st.expander("✅ Decisões Tomadas", expanded=False):
            st.success(structured.get("decisoes_tomadas", "Nenhuma decisão identificada."))

        with st.expander("📌 Plano de Ação Extraído", expanded=True):
            acoes = structured.get("acoes", [])
            if acoes:
                df_acoes = pd.DataFrame(acoes)
                col_map = {
                    "id": "ID", "acao": "Saída/Ação", "contexto": "Contexto",
                    "responsavel": "Responsável", "status": "Status", "prazo": "Prazo"
                }
                df_acoes = df_acoes.rename(columns={k: v for k, v in col_map.items() if k in df_acoes.columns})
                st.dataframe(df_acoes, use_container_width=True, hide_index=True)
            else:
                st.info("Nenhuma ação identificada.")

        # ── Download button ──
        def make_safe_filename(s: str) -> str:
            s = unicodedata.normalize("NFD", s).encode("ascii", "ignore").decode("ascii")
            return re.sub(r"[^\w\-]", "_", s).strip("_")

        safe_week = make_safe_filename(week)
        safe_type = make_safe_filename(meeting_type)
        filename = f"ATA_SOE_{safe_week}_{safe_type}.docx"

        st.markdown("---")
        col_d1, col_d2, col_d3 = st.columns([1, 2, 1])
        with col_d2:
            st.download_button(
                label="⬇️ Baixar Ata (.docx)",
                data=docx_bytes,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            st.caption(f"📁 Arquivo: {filename}")


if __name__ == "__main__":
    main()
